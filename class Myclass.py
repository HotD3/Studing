import datetime as DT
from datetime import timedelta
from ipaddress import collapse_addresses
from docx import Document
from docx.shared import Inches
document = Document()
#running = True
#date=input('Enter the date:')           #ввод текущей даты
#date1=DT.datetime.strptime(date, '%d/%m/%Y')
#print(date1)
today = DT.datetime.now()  #текущая дата
#timedelta=date1-today  #здесь должна вычисляться дельта. в данном случае возраст.
#print("{}-{}-{}".format(today.year, today.month, today.day))
#timedelta=(today - date1)
print('Date today:',today)

class Person:
    def print_info (self, n):
        for i in range(n):
            print(f"Name: {self.name}, Surname: {self.surname}, Where from: {self.Country}, Birth: {self.birth}, How_old: {self.old}")

p1 = Person()
p1.name = input('Enter name:')
p1.surname = input('Enter surname:')
p1.Country = input('Enter country of birth:')
p1.date= input('Enter date of birth:')
p1.birth = DT.datetime.strptime(p1.date, '%d/%m/%Y')
timedelta=p1.birth-today
p1.old = timedelta
p1.print_info(1)

p2 = Person()
p2.name = input('Enter name:')
p2.surname = input('Enter surname:')
p2.Country = input('Enter country of birth:')
p2.date= input('Enter date of birth:')
p2.birth = DT.datetime.strptime(p2.date, '%d/%m/%Y')
timedelta=p2.birth-today
p2.old = timedelta
p2.print_info(1)
records = (
    (p1.name , p1.surname , p1.Country , p1.date , p1.birth),
    (p2.name , p2.surname , p2.Country , p2.date , p2.birth)
)
document.add_heading('Таблица, в которую записываются данные, введенные в консоли',0)
table = document.add_table(rows = 1, cols=5)
col = table.rows[0].cells
col[0].text = 'Name'
col[1].text = 'Surname'
col[2].text = 'Where_from'
col[3].text = 'Birth'
col[4].text = 'How old'
for qw, e, r, t, y in records:
    row_cells = table.add_row().cells
    row_cells[0].text = qw
    row_cells[1].text = e
    row_cells[2].text = r
    row_cells[3].text = str(t)
    row_cells[4].text=  str(y)

document.save(r"C:\Users\Hot\Desktop\Code\document.docx")

#создать документ, куда будут экспортироваться личные данные

















#print(p1.old)
#print(now)
#while running:
 #   if y1<y2:
  #      print ('Correct')
   #     running = False
    #else:
     #   print('Date is not correct')
      #  running = False
        
        
   


2
#нужно узнать, сколько лет, имея дату рождения и текущуюб дату +
#сделать повторный ввод даты, если условие не выполняетс -
# нужно убрать отображение времени в текущей дате и дельте
